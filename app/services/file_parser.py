"""Multi-format file content extraction."""
import os
import csv
import json
import logging
from io import StringIO

logger = logging.getLogger(__name__)


async def parse_file(file_path: str) -> str:
    """Extract text content from a file. Returns the full text."""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()

        elif ext == ".csv":
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                reader = csv.DictReader(f)
                rows = []
                for row in reader:
                    rows.append(", ".join(f"{k}: {v}" for k, v in row.items()))
                return "\n".join(rows)

        elif ext == ".json":
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                return json.dumps(data, indent=2, default=str)

        elif ext == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                return text
            except ImportError:
                logger.warning("PyPDF2 not installed. Cannot parse PDF.")
                return "[PDF parsing requires PyPDF2. Install: pip install PyPDF2]"

        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(file_path)
                parts = []
                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            parts.append(row_text)
                return "\n".join(parts)
            except ImportError:
                logger.warning("python-docx not installed. Cannot parse DOCX.")
                return "[DOCX parsing requires python-docx. Install: pip install python-docx]"

        elif ext in (".xlsx", ".xls"):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True)
                text_parts = []
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    headers = []
                    for i, row in enumerate(ws.iter_rows(values_only=True)):
                        if i == 0:
                            headers = [str(c or "") for c in row]
                        else:
                            row_text = ", ".join(f"{headers[j] if j < len(headers) else f'col{j}'}: {str(c or '')}" for j, c in enumerate(row))
                            text_parts.append(row_text)
                return "\n".join(text_parts)
            except ImportError:
                logger.warning("openpyxl not installed. Cannot parse XLSX.")
                return "[XLSX parsing requires openpyxl. Install: pip install openpyxl]"

        else:
            return f"[Unsupported file type: {ext}]"

    except Exception as e:
        logger.error(f"File parsing error: {e}")
        return f"[Error parsing file: {str(e)}]"


def chunk_text(text: str, chunk_size: int = 1600, overlap: int = 320) -> list[str]:
    """Split text into overlapping chunks."""
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks
